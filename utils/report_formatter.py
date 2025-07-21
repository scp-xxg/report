"""
报告格式化工具
负责将生成的内容格式化为各种格式（Markdown、Word、PDF等）
"""
import json
import re
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches

class ReportFormatter:
    """
    报告格式化器
    
    设计理念：
    1. 多格式支持：支持输出多种格式的报告
    2. 模板化：使用模板系统便于定制报告样式
    3. 结构化：维护报告的层次结构和格式一致性
    4. 可扩展：易于添加新的输出格式
    """
    
    def __init__(self):
        self.report_data = {}
        self.creation_time = datetime.now()
    
    def format_markdown(self, report_data: Dict[str, Any]) -> str:
        """
        格式化为Markdown格式
        
        为什么选择Markdown：
        1. 简洁性：语法简单，易于编辑和阅读
        2. 通用性：被广泛支持，可以转换为多种格式
        3. 版本控制：纯文本格式，便于Git版本控制
        4. 协作性：便于团队协作和代码审查
        """
        markdown_content = []
        
        # 添加标题和元数据
        title = report_data.get('title', '未命名报告')
        markdown_content.append(f"# {title}\n")
        markdown_content.append(f"**生成时间**: {self.creation_time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        markdown_content.append(f"**生成系统**: 多智能体报告生成系统\n\n")
        
        # 添加目录
        if 'outline' in report_data:
            markdown_content.append("## 目录\n")
            for i, section in enumerate(report_data['outline'], 1):
                markdown_content.append(f"{i}. {section}\n")
            markdown_content.append("\n")
        
        # 添加各章节内容
        if 'sections' in report_data:
            for section_title, content in report_data['sections'].items():
                markdown_content.append(f"## {section_title}\n\n")
                markdown_content.append(f"{content}\n\n")
        
        # 添加图表引用
        if 'charts' in report_data:
            markdown_content.append("## 附录：图表\n\n")
            for chart_name, chart_path in report_data['charts'].items():
                markdown_content.append(f"### {chart_name}\n\n")
                markdown_content.append(f"![{chart_name}]({chart_path})\n\n")
        
        return "".join(markdown_content)
    
    def format_docx(self, report_data: Dict[str, Any], output_path: str):
        """
        格式化为Word文档
        
        为什么支持Word格式：
        1. 商业环境：Word是商业环境中最常用的文档格式
        2. 格式丰富：支持复杂的格式和布局
        3. 易于分享：便于与非技术人员分享
        4. 打印友好：适合打印和正式文档
        """
        doc = Document()
        
        # 添加标题
        title = report_data.get('title', '未命名报告')
        title_paragraph = doc.add_heading(title, level=0)
        
        # 添加元数据
        doc.add_paragraph(f"生成时间: {self.creation_time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("生成系统: 多智能体报告生成系统")
        doc.add_page_break()
        
        # 添加目录
        if 'outline' in report_data:
            doc.add_heading('目录', level=1)
            for i, section in enumerate(report_data['outline'], 1):
                doc.add_paragraph(f"{i}. {section}")
            doc.add_page_break()
        
        # 添加各章节内容
        if 'sections' in report_data:
            for section_title, content in report_data['sections'].items():
                doc.add_heading(section_title, level=1)
                # 处理段落
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
        
        # 保存文档
        doc.save(output_path)
        print(f"Word文档已保存到: {output_path}")
    
    def format_json(self, report_data: Dict[str, Any]) -> str:
        """
        格式化为JSON格式
        
        为什么需要JSON格式：
        1. 结构化：保持数据的完整结构信息
        2. 机器可读：便于程序处理和分析
        3. 接口友好：便于API接口返回
        4. 调试方便：便于开发和调试
        """
        # 添加元数据
        formatted_data = {
            'metadata': {
                'title': report_data.get('title', '未命名报告'),
                'creation_time': self.creation_time.isoformat(),
                'generator': '多智能体报告生成系统',
                'version': '1.0.0'
            },
            'content': report_data
        }
        
        return json.dumps(formatted_data, ensure_ascii=False, indent=2)
    
    def clean_text(self, text: str) -> str:
        """
        清理和标准化文本
        
        为什么需要文本清理：
        1. 一致性：确保格式的一致性
        2. 可读性：提高文本的可读性
        3. 兼容性：避免特殊字符导致的问题
        4. 质量：提升整体文档质量
        """
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 规范化换行符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 移除首尾空白
        text = text.strip()
        
        return text
    
    def add_table_of_contents(self, sections: List[str]) -> str:
        """
        生成目录
        
        设计思路：
        1. 自动编号：自动为章节添加编号
        2. 层次结构：支持多级目录结构
        3. 链接支持：在支持的格式中添加内部链接
        """
        toc_lines = ["## 目录\n"]
        for i, section in enumerate(sections, 1):
            toc_lines.append(f"{i}. {section}")
        
        return "\n".join(toc_lines) + "\n\n"
